"""
tools/docx_builder.py — Converts structured document content into a .docx file.
Handles headings, body paragraphs, inline images (by file path), and
an IEEE-formatted References section at the end.

Fixes applied:
  - All body text is justified (WD_ALIGN_PARAGRAPH.JUSTIFY)
  - Image width capped at 3.5 inches (centred)
  - Duplicate images skipped via MD5 hash tracking
  - Section heading space_before halved (6pt vs default 12-18pt)
"""
from __future__ import annotations
import hashlib
import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Set

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Regex to detect image path tokens embedded in section content
IMAGE_PATH_RE = re.compile(r"\[IMAGE_PATH:\s*(.+?)\]")
# Regex for any remaining unresolved placeholders (safety net)
PLACEHOLDER_RE = re.compile(r"\[IMAGE_PLACEHOLDER:[^\]]+\]")

# Max image width in inches — keeps images proportional and not oversized
IMAGE_MAX_WIDTH = Inches(3.5)


def build_docx(
    sections: List[Dict[str, Any]],
    references: List[str],
    output_path: str,
    title: str = "Generated Documentation",
) -> str:
    """
    Build a formatted .docx document.

    Args:
        sections: List of section dicts with 'title' and 'content' keys.
        references: Pre-formatted IEEE reference strings.
        output_path: Absolute path where the .docx should be saved.
        title: Document title for the cover heading.

    Returns:
        The output_path string on success.
    """
    doc = Document()
    _set_document_styles(doc)

    # Tracks MD5 hashes of inserted images to skip exact duplicates
    seen_image_hashes: Set[str] = set()

    # ── Cover Title ───────────────────────────────────────────────────────────
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_spacing(title_para, space_before=Pt(0), space_after=Pt(12))

    doc.add_paragraph()  # spacer

    # ── Body Sections ─────────────────────────────────────────────────────────
    for section in sections:
        heading_text = section.get("title", "Section")
        content = section.get("content", "")

        # Section heading (level 1) with reduced space before
        h = doc.add_heading(heading_text, level=1)
        _set_heading_spacing(h, space_before=Pt(6), space_after=Pt(4))

        # Write prose + images with deduplication
        _write_content_with_images(doc, content, seen_image_hashes)

    # ── References Section ────────────────────────────────────────────────────
    if references:
        ref_heading = doc.add_heading("References", level=1)
        _set_heading_spacing(ref_heading, space_before=Pt(6), space_after=Pt(4))
        for ref in references:
            p = doc.add_paragraph(ref, style="List Number")
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.save(output_path)
    logger.info(f"build_docx: saved → '{output_path}'")
    return output_path


# ── Internal Helpers ──────────────────────────────────────────────────────────

def _write_content_with_images(
    doc: Document, content: str, seen_hashes: Set[str]
) -> None:
    """
    Parse content text for [IMAGE_PATH: ...] tokens.
    Text blocks are written as justified paragraphs.
    Images are inserted centred, with duplicate detection via MD5 hash.
    """
    # Remove lingering unresolved placeholders
    content = PLACEHOLDER_RE.sub("", content)
    parts = IMAGE_PATH_RE.split(content)

    # parts alternates: [text, image_path, text, image_path, ..., text]
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Text block — write as justified paragraphs
            for line in part.split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(6)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            # Image path
            img_path = part.strip()
            _insert_image(doc, img_path, seen_hashes)


def _insert_image(doc: Document, img_path: str, seen_hashes: Set[str]) -> None:
    """
    Insert an image centred in the document.
    Skips if:
      - File does not exist
      - An identical image (same MD5 hash) was already inserted
    Width is capped at IMAGE_MAX_WIDTH to avoid oversized images.
    """
    path = Path(img_path)
    if not path.exists():
        logger.warning(f"_insert_image: file not found '{img_path}', skipping")
        return

    # Deduplication via MD5
    try:
        file_hash = hashlib.md5(path.read_bytes()).hexdigest()
    except OSError:
        file_hash = img_path  # Fallback to path string

    if file_hash in seen_hashes:
        logger.info(f"_insert_image: duplicate image skipped '{img_path}'")
        return
    seen_hashes.add(file_hash)

    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run()
        run.add_picture(img_path, width=IMAGE_MAX_WIDTH)
    except Exception as e:
        logger.warning(f"_insert_image: could not insert '{img_path}': {e}")


def _set_heading_spacing(paragraph, space_before: Pt, space_after: Pt) -> None:
    """Override space_before and space_after on a heading paragraph."""
    paragraph.paragraph_format.space_before = space_before
    paragraph.paragraph_format.space_after = space_after


def _set_document_styles(doc: Document) -> None:
    """Apply base font, margin, and paragraph style settings to the document."""
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    # Default paragraph style — Calibri 11pt, justified
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    style.paragraph_format.space_after = Pt(6)
